"""
Word document processor.
"""

from docx import Document
from documentor.config import logger
from documentor.processors.base import DocumentProcessor


class DocxProcessor(DocumentProcessor):
    """Processor for DOCX files"""
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text from a DOCX file
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text from the DOCX
        """
        try:
            doc = Document(file_path)
            paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text]
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        paragraphs.append(cell.text)
            
            return "\n".join(paragraphs)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
            return ""
